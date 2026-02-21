"""
Markdown → .docx 변환 스크립트
사용법: python scripts/md_to_docx.py <markdown_file> [output_file]

시나리오 Markdown 파일을 .docx로 변환합니다.
python-docx 패키지 필요: pip install python-docx

예시:
  python scripts/md_to_docx.py projects/kanta/output/final_screenplay.md
  python scripts/md_to_docx.py projects/kanta/output/final_screenplay.md output.docx
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("[!] python-docx가 설치되지 않았습니다.")
    print("    설치: pip install python-docx")
    sys.exit(1)


def create_screenplay_doc(md_path: Path, output_path: Path) -> None:
    """Markdown 시나리오를 .docx로 변환"""

    with open(md_path, encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # 페이지 설정 (A4, 한국 시나리오 여백)
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    # 스타일 정의
    styles = doc.styles

    # 씬 헤딩 스타일
    scene_style = styles.add_style("SceneHeading", 1)  # paragraph style
    scene_font = scene_style.font
    scene_font.name = "맑은 고딕"
    scene_font.size = Pt(11)
    scene_font.bold = True

    # 지문 스타일
    action_style = styles.add_style("Action", 1)
    action_font = action_style.font
    action_font.name = "맑은 고딕"
    action_font.size = Pt(10)

    # 인물명 스타일
    char_style = styles.add_style("CharacterName", 1)
    char_font = char_style.font
    char_font.name = "맑은 고딕"
    char_font.size = Pt(10)
    char_font.bold = True
    char_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 대사 스타일
    dialogue_style = styles.add_style("Dialogue", 1)
    dialogue_font = dialogue_style.font
    dialogue_font.name = "맑은 고딕"
    dialogue_font.size = Pt(10)
    dialogue_style.paragraph_format.left_indent = Cm(3.0)
    dialogue_style.paragraph_format.right_indent = Cm(2.0)

    # Markdown 파싱
    lines = content.split("\n")
    i = 0
    scene_heading_pattern = re.compile(r"^S#\d+\.")
    char_name_pattern = re.compile(r"^\s{10,}(\S+)")
    parenthetical_pattern = re.compile(r"^\s+\(.*\)")

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            # 빈 줄 → 단락 구분
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("#"):
            # Markdown 제목 → 무시 또는 제목으로
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("# ").strip()
            if level <= 2:
                doc.add_heading(text, level=min(level, 3))
            i += 1
            continue

        if scene_heading_pattern.match(stripped):
            # 씬 헤딩
            p = doc.add_paragraph(stripped, style="SceneHeading")
            i += 1
            continue

        if parenthetical_pattern.match(line):
            # 괄호 지시
            p = doc.add_paragraph(stripped, style="Dialogue")
            i += 1
            continue

        # 들여쓰기가 많으면 인물명 또는 대사로 판단
        leading_spaces = len(line) - len(line.lstrip())

        if leading_spaces >= 16:
            # 대사
            p = doc.add_paragraph(stripped, style="Dialogue")
        elif leading_spaces >= 10:
            # 인물명
            p = doc.add_paragraph(stripped, style="CharacterName")
        else:
            # 지문
            p = doc.add_paragraph(stripped, style="Action")

        i += 1

    doc.save(str(output_path))
    print(f"[+] .docx 저장 완료: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("사용법: python scripts/md_to_docx.py <markdown_file> [output_file]")
        sys.exit(1)

    md_file = Path(sys.argv[1])
    if not md_file.exists():
        print(f"[!] 파일을 찾을 수 없습니다: {md_file}")
        sys.exit(1)

    if len(sys.argv) >= 3:
        out_file = Path(sys.argv[2])
    else:
        out_file = md_file.with_suffix(".docx")

    create_screenplay_doc(md_file, out_file)
